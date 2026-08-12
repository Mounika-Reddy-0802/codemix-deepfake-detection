"""Render the mentor ethics sign-off form as a printable .docx.

The markdown at ``docs/ethics/mentor_signoff_form.md`` is the source of truth and
what gets reviewed in git; this script produces the copy the mentor actually signs
on paper. Re-run it whenever the markdown changes:

    python scripts/make_signoff_docx.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = "docs/ethics/mentor_signoff_form.docx"


def build() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.7)
        section.left_margin = section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    def heading(text: str, size: int = 13, space_before: int = 10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    def para(text: str, bold: bool = False, italic: bool = False, size: float = 10.5):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def table(headers, rows, widths=None, header: bool = True):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, text in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(9.5)
        for row in rows:
            cells = t.add_row().cells
            for i, text in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(str(text))
                run.font.size = Pt(9.5)
        if widths:
            for row in t.rows:
                for i, width in enumerate(widths):
                    row.cells[i].width = Inches(width)
        if not header:
            t.rows[0]._element.getparent().remove(t.rows[0]._element)
        return t

    # ---------------- Title ----------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ETHICS & DATA-USE SIGN-OFF")
    run.bold = True
    run.font.size = Pt(17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty Mentor Approval")
    run.font.size = Pt(12.5)

    table(
        ["Field", "Detail"],
        [
            [
                "Project",
                "Quantifying and Closing the Code-Mixing Generalisation Gap in Audio "
                "Deepfake Detection, with a Live Call-Detection Demonstrator",
            ],
            ["Programme", "B.Tech CSE (Data Science), Semester 7 - Capstone Project"],
            [
                "Team",
                "M. Lahari (D006) | S. Mounika Reddy (D032) | K. Sai Krishna Reddy (D034)",
            ],
            ["Repository", "github.com/Mounika-Reddy-0802/codemix-deepfake-detection"],
            ["Approval sought for", "The REVISED (v2) project scope, dated 12 August 2026"],
        ],
        widths=[1.5, 5.4],
    )

    # ---------------- 1 ----------------
    heading("1.  What the project does")
    para(
        "We are building a DETECTOR that identifies synthetic (AI-cloned) speech in "
        "Hindi-English code-mixed telephone calls - the kind used in voice-based fraud in "
        "India. Existing detectors are trained almost entirely on English and collapse on "
        "code-mixed telephony audio; we measure that failure, reduce it, and demonstrate it "
        "on a live call. To train and honestly test such a detector, the project must itself "
        "generate synthetic speech. That generation is the ethically sensitive activity, and "
        "it is what this form seeks approval for."
    )
    para(
        "No synthetic speech of any kind has been generated to date. The code that would "
        "generate it is written but is blocked in software until this form is signed and "
        "uploaded (see Section 5).",
        bold=True,
    )

    # ---------------- 2 ----------------
    heading("2.  What changed since the original plan - this is what needs fresh approval")
    table(
        ["Item", "Originally approved", "Revised (v2) scope"],
        [
            ["Volume of generated speech", "~1,500-2,500 clips", "~4,000+ clips"],
            [
                "Attack families",
                "XTTS-v2 (TTS) + Tortoise (held-out)",
                "+ RVC voice conversion (10-15 speaker models)",
            ],
            [
                "External evaluation",
                "None",
                "AffectDF public benchmark subset (CC BY-NC 4.0)",
            ],
            ["Everything else", "-", "Unchanged"],
        ],
        widths=[1.7, 2.3, 2.9],
    )
    para(
        "Scientific reason for the increase: a detector trained on too little synthetic "
        "speech, or on a single generator, learns the quirks of that one tool rather than the "
        "general signature of synthesis. More volume and a second, structurally different "
        "attack family (voice conversion vs text-to-speech) is what makes the result "
        "trustworthy rather than an artefact."
    )

    # ---------------- 3 ----------------
    heading("3.  Data sources and their licences")
    para(
        "All corpora are public research datasets. No data is collected from any individual "
        "by this project."
    )
    table(
        ["Corpus", "Licence", "How we use it"],
        [
            ["ASVspoof 2019 LA", "Research (ASVspoof)", "Training the English baseline"],
            [
                "MUCS 2021 Hindi-English (OpenSLR 104)",
                "CC BY-SA 4.0",
                "Code-mixed speech; voice-cloning references",
            ],
            ["HiACC - ADULT SUBSET ONLY", "CC BY 4.0", "Code-mixed evaluation speech"],
            ["IndicVoices (AI4Bharat)", "Research, gated", "Hindi / Tamil evaluation"],
            ["IndicSynth", "CC BY-NC 4.0", "Evaluation only"],
            ["IndicTTS-Deepfake", "Research", "Evaluation only"],
            ["AffectDF", "CC BY-NC 4.0", "External evaluation only"],
        ],
        widths=[2.5, 1.6, 2.8],
    )
    para(
        "Generation tools: Coqui XTTS-v2 (CPML - non-commercial), Tortoise-TTS (Apache-2.0), "
        "RVC (open source)."
    )
    para(
        "Two licences - IndicSynth and XTTS-v2 - permit NON-COMMERCIAL use only. This project "
        "is non-commercial academic research, which both permit. The team commits to keeping "
        "it so."
    )

    doc.add_page_break()

    # ---------------- 4 ----------------
    heading("4.  The rules the team commits to", space_before=0)
    rules = [
        (
            "4.1  HiACC child audio is excluded entirely. This is absolute.",
            "The HiACC corpus contains 1,858 utterances from CHILDREN. These are never used "
            "for anything: not as genuine speech, not as a voice-cloning reference, not in any "
            "data listing. They are quarantined into a separate folder at download time and "
            "the pipeline refuses to read it.",
        ),
        (
            "4.2  Voice cloning is limited to adult speakers in licensed research corpora,",
            "solely to build detection training data.",
        ),
        (
            "4.3  No real, identifiable private individual will be cloned.",
            "No team member's voice, no third party's voice, no public figure's voice - not "
            "for testing, not for the demonstration.",
        ),
        (
            "4.4  Generated clones are never released as a standalone clone set.",
            "Any release of derived data respects the source licence (MUCS is share-alike). "
            "The detector code and trained model may be released; the raw synthetic speech "
            "will not be published in a form usable as a cloning resource.",
        ),
        ("4.5  No attempt will be made to re-identify any speaker in any corpus.", ""),
        (
            "4.6  The demonstration is defensive.",
            "The live call demo warns the person RECEIVING the call. It never signals the "
            "caller, never disconnects a call automatically, and produces no output that would "
            "help someone commit fraud.",
        ),
        ("4.7  Non-commercial academic use only, as required by the licences above.", ""),
        (
            "4.8  Scientific integrity.",
            "The English baseline is trained on ASVspoof 2019 only; evaluation-only corpora "
            "and the held-out attack tool never enter training data. This is what makes the "
            "paper's central claim measurable rather than circular.",
        ),
        (
            "4.9  Generated audio and credentials stay on team-controlled storage",
            "and are never committed to the public repository.",
        ),
        (
            "4.10  Every corpus and tool will be cited as its licence requires,",
            "and the licence restrictions above will be stated explicitly in the published paper.",
        ),
    ]
    for title, body in rules:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        if body:
            run = p.add_run(" " + body)
            run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Disclosure, in the interest of honesty: ")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run = p.add_run(
        "on 12 August 2026 the team found that although the child folder was being "
        "quarantined, a later processing step would still have read it. No child audio was "
        "ever actually processed - the corpus has not been downloaded yet. The defect was "
        "fixed, an audit tool was written, and automated tests now make the exclusion "
        "impossible to reopen silently. We report this because a rule that is only checked by "
        "memory is not a rule."
    )
    run.italic = True
    run.font.size = Pt(10)

    # ---------------- 5 ----------------
    heading("5.  How these rules are enforced in software, not just promised")
    table(
        ["Rule", "Enforcement"],
        [
            [
                "No generation before this signature",
                "An ETHICS GATE in the code refuses to load any voice-cloning model until a "
                "signed copy of this form exists in the repository. It has NO override switch "
                "and cannot be bypassed by a setting.",
            ],
            [
                "Child-audio exclusion",
                "Automated tests plus an audit tool that reports any child-looking file and "
                "refuses to declare the check passed without a human signature.",
            ],
            [
                "No leakage between training and testing",
                "An automated test suite that runs before every training job.",
            ],
            [
                "Speaker groups fixed in advance",
                "Speaker lists are frozen and fingerprinted (SHA-256); any later change is "
                "detected.",
            ],
        ],
        widths=[2.0, 4.9],
    )

    # ---------------- 6 ----------------
    heading("6.  Declaration by the student team")
    para(
        "We have read the rules in Section 4 and will abide by them for the duration of the "
        "project. We understand that generating synthetic speech carries a risk of misuse and "
        "that our obligation is to keep this work defensive, non-commercial, and properly "
        "attributed."
    )
    t = table(
        ["Name", "Roll No.", "Signature", "Date"],
        [
            ["M. Lahari", "D006", "", ""],
            ["S. Mounika Reddy", "D032", "", ""],
            ["K. Sai Krishna Reddy", "D034", "", ""],
        ],
        widths=[2.0, 0.9, 2.6, 1.4],
    )
    for row in t.rows[1:]:
        row.height = Inches(0.45)

    # ---------------- 7 ----------------
    heading("7.  Faculty mentor approval")
    para(
        "I have reviewed the revised scope in Section 2, the data sources and licences in "
        "Section 3, and the commitments in Section 4. I approve the team to proceed with the "
        "generation of synthetic speech for the purpose of building and evaluating a "
        "detector, subject to any conditions noted below."
    )
    p = doc.add_paragraph()
    run = p.add_run("Decision (tick one):     ")
    run.bold = True
    run = p.add_run(
        "[  ] Approved          [  ] Approved with conditions below          [  ] Not approved"
    )
    run.font.size = Pt(11)

    para("Conditions / remarks:", bold=True)
    for _ in range(3):
        doc.add_paragraph("_" * 95)

    t = table(
        ["", ""],
        [
            ["Name", ""],
            ["Designation", ""],
            ["Department", ""],
            ["Signature", ""],
            ["Date", ""],
        ],
        widths=[1.5, 5.4],
        header=False,
    )
    for row in t.rows:
        row.height = Inches(0.38)
        row.cells[0].paragraphs[0].runs[0].bold = True

    para("")
    para("Institution stamp (if required):", italic=True)
    for _ in range(4):
        doc.add_paragraph("")
    para(
        "This form and the supporting licence register (docs/licences.md) are maintained in "
        "the project repository. The signed copy is scanned and stored at "
        "docs/ethics/mentor_signoff_<yyyy-mm-dd>.pdf",
        italic=True,
        size=8.5,
    )
    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"wrote {OUT} ({os.path.getsize(OUT) / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
